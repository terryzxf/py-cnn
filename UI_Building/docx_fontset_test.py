from docx import Document
#import docx
#from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 引入文字的位置样式
from docx.shared import Pt  # 设置字体
from docx import Document
filepath = '/home/zxf/病人登记/记录/肖文德_放疗前小结.doc'

ad = Document(filepath)
doctext=[paragraph.text for paragraph in ad.paragraphs]

#print(type(doctext))

#for i in range(len(doctext)):
    #print ('**'+str(i)+'*\n',doctext[i])

#print(doctext[4].find('R'))
#print(doctext[7].find('R'))
#print(doctext[9].find('R'))

#4

p = ad.paragraphs[4].clear()
run0 = p.add_run(doctext[4][:8])
font0 = run0.font
font0.size = Pt(11)
font0.bold = True


#if not doctext[4].find('R') == 8:
run1 = p.add_run(doctext[4][8:doctext[4].find('R')])
font = run1.font
font.size = Pt(11)

run2 = p.add_run('R')
font2 = run2.font
font2.size = Pt(11)
font2.name = 'wingdings 2'

run3 = p.add_run(doctext[4][doctext[4].find('R')+1:])
font3 = run3.font
font3.size = Pt(11)

#7
p = ad.paragraphs[7].clear()
run0 = p.add_run(doctext[7][:6])
font0 = run0.font
font0.size = Pt(11)
font0.bold = True


#if doctext[7].find('R') == 6:
run1 = p.add_run(doctext[7][6:doctext[7].find('R')])
font = run1.font
font.size = Pt(11)

run2 = p.add_run('R')
font2 = run2.font
font2.size = Pt(11)
font2.name = 'wingdings 2'

run3 = p.add_run(doctext[7][doctext[7].find('R')+1:])
font3 = run3.font
font3.size = Pt(11)

#8

p = ad.paragraphs[8].clear()
run0 = p.add_run(doctext[8][:6])
font0 = run0.font
font0.size = Pt(11)
font0.bold = True

#if not doctext[8].find('R') == 6:
run1 = p.add_run(doctext[8][6:doctext[8].find('R')])
font = run1.font
font.size = Pt(11)

run2 = p.add_run('R')
font2 = run2.font
font2.size = Pt(11)
font2.name = 'wingdings 2'

run3 = p.add_run(doctext[8][doctext[8].find('R')+1:])
font3 = run3.font
font3.size = Pt(11)

#9
p = ad.paragraphs[9].clear()
run0 = p.add_run(doctext[9][:8])
font0 = run0.font
font0.size = Pt(11)
font0.bold = True

#if not doctext[9].find('R') == 8:
run1 = p.add_run(doctext[9][8:doctext[9].find('R')])
font = run1.font
font.size = Pt(11)

run2 = p.add_run('R')
font2 = run2.font
font2.size = Pt(11)
font2.name = 'wingdings 2'

run3 = p.add_run(doctext[9][doctext[9].find('R')+1:])
font3 = run3.font
font3.size = Pt(11)

ad.save(filepath)

#4 7 8 9
"""
print(doctext[4][:doctext[4].find('R')])
print(doctext[4][doctext[4].find('R')])
print(doctext[4][doctext[4].find('R')+1:])

print(doctext[7][:doctext[7].find('R')])
print(doctext[7][doctext[7].find('R')])
print(doctext[7][doctext[7].find('R')+1:])

print(doctext[8][:doctext[8].find('R')])
print(doctext[8][doctext[8].find('R')])
print(doctext[8][doctext[8].find('R')+1:])

print(doctext[9][:doctext[9].find('R')])
print(doctext[9][doctext[9].find('R')])
print(doctext[9][doctext[9].find('R')+1:])

"""