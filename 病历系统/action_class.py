
# from date_calculate import calculate_end_date
from docxtpl import DocxTemplate
from docx import Document
import arrow as arw
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 引入文字的位置样式
from docx.shared import Pt  # 设置字体
import sys


gerpath = str(sys.path[0])


# 计算并返回 开始日期间隔次数后的结束日期
def calculate_end_date(startday, dosesplit):
    st = startday  # 开始日期
    T = dosesplit  # 次数
    # test tail
    # T = '29'
    # st = '2018-02-01'

    # test tail
    # print(type(weekend_judge))
    # print ('  weekend_judge:',weekend_judge)
    # print('  T_judge:',T_judge)
    # print('  eval(T):',eval(T))
    # print('  not(isinstance(eval(T),int)):',not(isinstance(eval(T),int)))

    st_date = arw.get(st)
    a = int(T) // 5
    b = int(T) % 5
    # test tail
    # print('   ', st_date.isoweekday())
    if b == 0:
        if st_date.isoweekday() == 1:
            day_shift = (a - 1) * 7 + 5
        else:
            day_shift = (a - 1) * 7 + 5 + 2

    elif b == 1:
        day_shift = a * 7 + b

    elif b == 2:
        if st_date.isoweekday() == 5:
            day_shift = a * 7 + b + 2
        else:
            day_shift = a * 7 + b

    elif b == 3:
        if st_date.isoweekday() == 4 or st_date.isoweekday() == 5:
            day_shift = a * 7 + b + 2
        else:
            day_shift = a * 7 + b

    elif b == 4:
        if st_date.isoweekday() == 1 or st_date.isoweekday() == 2:
            day_shift = a * 7 + b
        else:
            day_shift = a * 7 + b + 2
    end_date = st_date.shift(days=day_shift)
    out_enddate = end_date.format('YYYY-MM-DD')
    return out_enddate


# 计算首次计划剂量模式输入Tdose(xx/ff):   返回tdose('xxGy/yyF'), 分割次数, 单次剂量
def dose_cal(Tdose):

    dose = Tdose.split('/')   # 对TdoseE输入格式 xxGY/yyF 的处理 分割成 [xx,yy]
    tdose = dose[0]+'Gy/'+dose[1]+'f'      # 首次计划的总剂量
    dosesplit = eval(dose[1]) # 首次计划的次数
    dose_t = eval(dose[0]) / dosesplit  # 单次剂量
    return tdose, dosesplit, dose_t   # 返回首次总剂量, 分割次数, 单次剂量


# 计算追加剂量信息追加剂量的输入格式(剂量/次数/开始日期(yyyy-mm-dd)/野数 )     返回tdose(xxGy/yyF)  分割次数 单次剂量 开始日期 野数
def EXLdose_cal(Tdose):

    dose = Tdose.split('/')   # 对TdoseE输入格式 xxGY/yyF 的处理 分割成 [xx,yy]
    tdose = dose[0]+'Gy/'+dose[1]+'f'      # 首次计划的总剂量
    dosesplit = eval(dose[1]) # 首次计划的次数
    dose_t = eval(dose[0])/dosesplit  # 单次剂量
    startDate = dose[2]
    ExbeamNum = dose[3]
    return tdose, dosesplit, dose_t, startDate, ExbeamNum  # 返回首次总剂量, 分割次数, 单次剂量


# 计算各个病程的日期 返回数组date
def date_cal(startdate, dosesplit): # 计算出每个病程的日期存入列表

    date = []
    # enddate = calculate_end_date(startdate, dosesplit)  #结束日期
    recNum = int(dosesplit/5+1)
    # print(recNum,startdate)

    for i in range(recNum):

        if i == 0:

            date.append(startdate)   #初始日期
        else:
            date.append(calculate_end_date(startdate, 5 * i-1)) # 每5次的日期
    return date      # 返回日期列表 列表长度


#生成病程
def doc_mk(pt_S, beamNum, firmStyle, Tdose, startdate, diagnosisT, sectNum, Exdose, Ldose, doctor='赵晓峰'):

        #    需要传入:  病人信息  治疗野数^  固定方式^  医师  总治疗剂量^  date[] 特殊表现({{diagnosisT}})   需要UI支持三个量输入
        #   1.标题生成  done
        #   2.首次病程生成  done
        #   3.计算需要的中间病程数  ing...  计算结束日期OK , 中间记录数= 分割次数/5-2  计算每次日期. date[]
        #   4.生成自动的中间病程    循环实现  特殊表现代入与循环变量挂钩  5, 10, 15, 20, 25, 30, 35,特殊表现( {{ diagnosisT[] }} )
        #   5.文档变量代入  模板变量代入
        #   6.文档保存 文件名 病人姓名_病程.docx
        #   文档实现功能: 文档生成病程记录: 自动实现各个信息变量代入,格式代入,自动保存.

        # info_get()返回 pt_S 为病人信息序列: ['id', 'RTid','Curesta', 'Names', 'sex', 'age', 'Department',
        #                                   'bedNum','dignose', 'phoneNum', 'bodyPart', 'startDate']
        #
        #  ['姓名:', pt_S.Names,       '性别:', pt_S.sex,           '年龄:', str(pt_S.age),
        #   '科室:', pt_S.Department,  '床号:', str(pt_S.bedNum),
        #   '诊断:', pt_S.dignose,     '电话:', str(pt_S.phoneNum), ,
        #   '治疗部位:', str(pt_S.bodyPart),    '开始日期:', str(pt_S.startDate)]
        #
        #   模板文档; 首次病程: X1doc.docx  段落文档: paragraph.docx
        #   需迭代变量:
        # {{startDate}}  --  date[0]^^^             date[]数组构建
        # {{dateN}}          date[n]
        # {{dateL}}          date[L]
        # {{pName}},         pt_S.Names
        # {{sex}},           pt_S.sex
        # {{age}}            pt_S.age
        # {{diagnose}}       pt_S.dignose
        # {{bodyPart}}       pt_S.bodyPart
        # {{confirm}}        firmStyle              固定方式传倒
        # {{beamNum}}        beamNum                野数传到
        # {{Dose}}           Tdose                  剂量推倒分割方式推倒 dose[]构建
        # {{DoseN}}          dose[n]
        # {{DoseL}}          Tdose/ Tdose+Ldose[]   分段模式构建总剂量计算
        # {{doctor}}    --   默认值
        # 内在变量sectNum 治疗分段改野次数
        # 函数传入列表 pt_S, beamNum, firmStyle, Tdose, date, diagnosisT Ldose, sectNum)

        # UI需要构建传入结构: 治疗野数^  固定方式^  总治疗剂量^  开始治疗日期^  分段数  [ Ldose[] ]
        # 野数输入框  固定方式下拉菜单  总剂量输入  开始日期输入框 分段数   根据分段数变化的LDose[]输入框数


    gerpath = str(sys.path[0])
    tdose, dosesplit, dose_t = dose_cal(Tdose)
    fdate = date_cal(startdate,dosesplit)
    Rectimes = len(fdate)

    file1 = docx.Document(gerpath+'/模板doc/X1doc.docx')
    file2 = docx.Document(gerpath+'/模板doc/paragraph.docx')
    docpath =gerpath+'/record/病程记录/' + pt_S.Names + '_病程记录.docx'
    signment = gerpath+ '/模板doc/3.png'
    doc = docx.Document() # 'home/zxf/病人登记/病程记录/'+ pt_S.Names +'_病程记录.docx'  # 新建word文档
    docTex1 = [paragraph.text for paragraph in file1.paragraphs]
    # file1[0] = file1[0].replace('{{pName}}', pt_S.Names)

    docTex1[0] = docTex1[0].format(pName=pt_S.Names, sex=pt_S.sex, age=pt_S.age, diagnose=pt_S.dignose, bodyPart=pt_S.bodyPart)
    docTex1[1] = docTex1[1].format(confirm=firmStyle, bodyPart=pt_S.bodyPart, beamNum= beamNum, Dose = tdose)
    docTex2 = [paragraph.text for paragraph in file2.paragraphs]

    diaT = ['', '，受照区无色素沉着，表皮完整无破溃', '，受照区轻度色素沉着，表皮完整无破溃', '，受照区轻度色素沉着，表皮完整无破溃',
            '，受照区轻度色素沉着，轻度发红，表皮完整无破溃','，受照区中度色素沉着，表皮完整无破溃',
            '，受照区中重度色素沉着，局部皮肤小范围干性脱皮，表皮完整无破溃','，受照区重度色素沉着，局部皮肤干性脱皮明显，无渗出，表皮完整无破溃']

    paragraph1 = doc.add_paragraph()
    # paragraph1.paragragh_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = paragraph1.add_run('病程记录')

    # title = parTitle.add_run()
    title.font.size = Pt(16)
    title.font.bold = True
    # date=['20180101','20180107','20180201','20180301','20180401']

    for n in range(Rectimes):

        if n == 0:

            firstPh = doc.add_paragraph()
            firstPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            ad0 = firstPh.add_run(fdate[n])
            ad0.font.size = Pt(11)

            for j in range(len(docTex1)):

                firstPh = doc.add_paragraph()
                firstPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                ad0 = firstPh.add_run('      '+ docTex1[j])
                ad0.font.size = Pt(11)

            firsts = doc.add_paragraph()
            firsts.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            ad0s = firsts.add_run()
            ad0s.add_picture(signment)
            ad0s.font.size = Pt(11)

        elif n>0 and n <Rectimes-1:

            midPhdoc = docTex2[0].format(bodyPart=pt_S.bodyPart, n=n*5, doseS=5*n*dose_t, diagnosisT=diaT[n])
            midD = doc.add_paragraph()
            midD.paragraph_format.alignment =WD_PARAGRAPH_ALIGNMENT.LEFT
            ad1D = midD.add_run(fdate[n])
            ad1D.font.size = Pt(11)

            midPh = doc.add_paragraph()
            midPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            ad1 = midPh.add_run('      '+ midPhdoc)
            ad1.font.size = Pt(11)

            mids = doc.add_paragraph()
            mids.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            ad1s = mids.add_run()
            ad1s.add_picture(signment)
            ad1s.font.size = Pt(11)

        else:
            if sectNum == '':
                endsay = '今日放疗结束，嘱其放疗后注意事项，建议酌情继续行抗肿瘤综合治疗，'
            else:
                endsay = ''
            endDdoc = docTex2[1].format(bodyPart=pt_S.bodyPart, n=dosesplit, doseL= 5*n*dose_t, diagnosisT=diaT[n], endsay=endsay)

            endD = doc.add_paragraph()
            endD.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            ad2D = endD.add_run(fdate[n])
            ad2D.font.size = Pt(11)
            endPh = doc.add_paragraph()
            endPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # ad2 = endPh.add_run(date[n] + '\n')
            ad2 = endPh.add_run('      '+endDdoc)
            ad2.font.size = Pt(11)

            ends = doc.add_paragraph()
            ends.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            ad2s = ends.add_run()
            ad2s.add_picture(signment)

            ad2s.font.size = Pt(11)

    # 第一次补量的记录
    if sectNum != '' and Exdose!= '':
        Extdose, Exsplit, Exdose_t, ExstaDate, ExbeamNum = EXLdose_cal(Exdose)
        dateEx= date_cal(ExstaDate,Exsplit)
        ExRectimes = len(dateEx)

        for n in range(ExRectimes):

            if n == 0:

                firstPh = doc.add_paragraph()
                firstPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                ad0 = firstPh.add_run(dateEx[n])
                ad0.font.size = Pt(11)

                Exfstdoc = docTex2[2].format(bodyPart=pt_S.bodyPart, n=dosesplit, doseL=dosesplit * dose_t,
                                             diagnosisT=diaT[int(dosesplit/5 + n)], confirm=firmStyle,
                                             beamNum=ExbeamNum, Dose=Extdose)
                firstPh = doc.add_paragraph()
                firstPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                ad0 = firstPh.add_run('      ' + Exfstdoc)
                ad0.font.size = Pt(11)

                firsts = doc.add_paragraph()
                firsts.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                ad0s = firsts.add_run()
                ad0s.add_picture(signment)
                ad0s.font.size = Pt(11)

            elif n > 0 and n < ExRectimes - 1 and ExRectimes>2:

                midPhdoc = docTex2[0].format(bodyPart=pt_S.bodyPart, n=n * 5+dosesplit ,
                                             doseS=5 * n * Exdose_t+ dosesplit*dose_t, diagnosisT=diaT[int(dosesplit/5 + n)])
                midD = doc.add_paragraph()
                midD.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                ad1D = midD.add_run(dateEx[n])
                ad1D.font.size = Pt(11)

                midPh = doc.add_paragraph()
                midPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                ad1 = midPh.add_run('      ' + midPhdoc)
                ad1.font.size = Pt(11)

                mids = doc.add_paragraph()
                mids.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                ad1s = mids.add_run()
                ad1s.add_picture(signment)
                ad1s.font.size = Pt(11)

            else:
                if sectNum == '1':
                    endsay = '今日放疗结束，嘱其放疗后注意事项，建议酌情继续行抗肿瘤综合治疗，'
                else:
                    endsay = ''
                endDdoc = docTex2[1].format(bodyPart=pt_S.bodyPart, n=dosesplit + Exsplit,
                                            doseL=dosesplit * dose_t + Exsplit * Exdose_t,
                                            diagnosisT=diaT[n], endsay=endsay)

                endD = doc.add_paragraph()
                endD.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                ad2D = endD.add_run(dateEx[n])
                ad2D.font.size = Pt(11)

                endPh = doc.add_paragraph()
                endPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                # ad2 = endPh.add_run(date[n] + '\n')
                ad2 = endPh.add_run('      ' + endDdoc)
                ad2.font.size = Pt(11)

                ends = doc.add_paragraph()
                ends.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                ad2s = ends.add_run()
                ad2s.add_picture(signment)
                ad2s.font.size = Pt(11)

    # 第二次补量纪录
    if sectNum >='1' and Ldose != '':
        Ltdose, Lsplit, Ldose_t, LstaDate, LbeamNum = EXLdose_cal(Ldose)
        dateL = date_cal(LstaDate, Lsplit)
        LRectimes = len(dateL)

        for n in range(LRectimes):
            print('#',n,'#\n' )
            if n == 0:

                firstPh = doc.add_paragraph()
                firstPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                ad0 = firstPh.add_run(dateL[n])
                ad0.font.size = Pt(11)

                Exfstdoc = docTex2[2].format(bodyPart=pt_S.bodyPart, n=dosesplit+Exsplit,
                                             doseL=dosesplit * dose_t+Exsplit*Exdose_t,
                                             diagnosisT=diaT[int(dosesplit / 5 + n)],
                                             confirm=firmStyle,
                                             beamNum=LbeamNum, Dose=Ltdose)
                firstPh = doc.add_paragraph()
                firstPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                ad0 = firstPh.add_run('      ' + Exfstdoc)
                ad0.font.size = Pt(11)

                firsts = doc.add_paragraph()
                firsts.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                ad0s = firsts.add_run()
                ad0s.add_picture(signment)
                ad0s.font.size = Pt(11)

            elif n > 0 and n < ExRectimes - 1 and LRectimes>2:

                midPhdoc = docTex2[0].format(bodyPart=pt_S.bodyPart, n=n * 5+dosesplit+Exsplit,
                                             doseS=5 * n * Ldose_t +dosesplit*dose_t+Exsplit*Exdose_t,
                                             diagnosisT=diaT[int(dosesplit / 5 + n)])
                midD = doc.add_paragraph()
                midD.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                ad1D = midD.add_run(dateL[n])
                ad1D.font.size = Pt(11)

                midPh = doc.add_paragraph()
                midPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                ad1 = midPh.add_run('      ' + midPhdoc)
                ad1.font.size = Pt(11)

                mids = doc.add_paragraph()
                mids.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                ad1s = mids.add_run()
                ad1s.add_picture(signment)
                ad1s.font.size = Pt(11)

            else:
                if sectNum == '2':
                    endsay = '今日放疗结束，嘱其放疗后注意事项，建议酌情继续行抗肿瘤综合治疗，'
                else:
                    endsay = ''
                endDdoc = docTex2[1].format(bodyPart=pt_S.bodyPart, n=dosesplit+n*5+Exsplit,
                                            doseL=5 * n * Ldose_t +dosesplit*dose_t+Exsplit*Exdose_t,
                                            diagnosisT=diaT[n], endsay=endsay)

                endD = doc.add_paragraph()
                endD.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                print(n,dateL,LRectimes)

                ad2D = endD.add_run(dateL[n])
                ad2D.font.size = Pt(11)

                endPh = doc.add_paragraph()
                endPh.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                # ad2 = endPh.add_run(date[n] + '\n')
                ad2 = endPh.add_run('      ' + endDdoc)
                ad2.font.size = Pt(11)

                ends = doc.add_paragraph()
                ends.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                ad2s = ends.add_run()
                ad2s.add_picture(signment)
                ad2s.font.size = Pt(11)

    doc.save(docpath)


"""
pt_S, patient_info, table_tail = info_get.info_get(3000)
beamNum = 3
firmStyle = 'U型面膜+体架'
Tdose= '40/20'
date=['20180101','20180107','20180201','20180301','2018041']
diagnosisT = '@@@@@'
sectNum ='2'
Exdose ='20/10/2018-05-01/4'
Ldose ='10/5/2018-06-01/5'
doc_mk(pt_S, beamNum, firmStyle, Tdose, date, diagnosisT, sectNum, Exdose, Ldose)

"""
"""
#第一段
paragraph1=file.add_paragraph('病程记录')
paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
paragraph1.style='Title'

#第二段
paragraph2=file.add_paragraph()
paragraph2.paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
L = 'text'
run=paragraph2.add_run(L) #向段落中追加文字
run.bold=True    #设置追加文字样式
run.font.size=Pt(30)
#run.style='Emphasis'

# 第三段
paragraph3=file.add_paragraph()
paragraph3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
paragraph3.style = 'Normal'
run=paragraph3.add_run('Tomorrow is my birthay.I am looking forword your coming')
run.font.size=Pt(20)
file.save('/home/zxf/病人登记/模板doc/text.docx')
#file.save('/home/zxf/病人登记/模板doc/'+ L +'.docx')
"""


    # 放疗前小结(放疗治疗计划小结的格式修改)
def ad_rec_style(filepath):
    # from docx import Document
    # import docx
    # from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 引入文字的位置样式

    # filepath = '/home/zxf/病人登记/记录/肖文德_放疗前小结.doc'

    ad = Document(filepath)
    doctext = [paragraph.text for paragraph in ad.paragraphs]

    # print(type(doctext))

    # for i in range(len(doctext)):
      #  print ('**'+str(i)+'*\n',doctext[i])

    # print(doctext[4].find('R'))
    # print(doctext[7].find('R'))
    # print(doctext[9].find('R'))


    # 4

    p = ad.paragraphs[4].clear()
    run0 = p.add_run(doctext[4][:8])
    font0 = run0.font
    font0.size = Pt(11)
    font0.bold = True

    # if not doctext[4].find('R') == 8:
    run1 = p.add_run(doctext[4][8:doctext[4].find('R')])
    font = run1.font
    font.size = Pt(11)

    run2 = p.add_run('R')
    font2 = run2.font
    font2.size = Pt(11)
    font2.name = 'wingdings 2'

    run3 = p.add_run(doctext[4][doctext[4].find('R') + 1:])
    font3 = run3.font
    font3.size = Pt(11)

    # 7
    p = ad.paragraphs[7].clear()
    run0 = p.add_run(doctext[7][:6])
    font0 = run0.font
    font0.size = Pt(11)
    font0.bold = True

    # if doctext[7].find('R') == 6:
    run1 = p.add_run(doctext[7][6:doctext[7].find('R')])
    font = run1.font
    font.size = Pt(11)

    run2 = p.add_run('R')
    font2 = run2.font
    font2.size = Pt(11)
    font2.name = 'wingdings 2'

    run3 = p.add_run(doctext[7][doctext[7].find('R') + 1:])
    font3 = run3.font
    font3.size = Pt(11)

    # 8

    p = ad.paragraphs[8].clear()
    run0 = p.add_run(doctext[8][:6])
    font0 = run0.font
    font0.size = Pt(11)
    font0.bold = True

    # if not doctext[8].find('R') == 6:
    run1 = p.add_run(doctext[8][6:doctext[8].find('R')])
    font = run1.font
    font.size = Pt(11)

    run2 = p.add_run('R')
    font2 = run2.font
    font2.size = Pt(11)
    font2.name = 'wingdings 2'

    run3 = p.add_run(doctext[8][doctext[8].find('R') + 1:])
    font3 = run3.font
    font3.size = Pt(11)

    # 9
    p = ad.paragraphs[9].clear()
    run0 = p.add_run(doctext[9][:8])
    font0 = run0.font
    font0.size = Pt(11)
    font0.bold = True

    # if not doctext[9].find('R') == 8:
    run1 = p.add_run(doctext[9][8:doctext[9].find('R')])
    font = run1.font
    font.size = Pt(11)

    run2 = p.add_run('R')
    font2 = run2.font
    font2.size = Pt(11)
    font2.name = 'wingdings 2'

    run3 = p.add_run(doctext[9][doctext[9].find('R') + 1:])
    font3 = run3.font
    font3.size = Pt(11)
    # 11  28 31

    ad.save(filepath)


def rec_mk(pt_S, startDate, rtid):
    context = dict(name=pt_S.Names, sex=pt_S.sex, age=pt_S.age,
                   nid=pt_S.id, dignose=pt_S.dignose, part=pt_S.bodyPart,
                   date=startDate, bedNum=pt_S.bedNum, department=pt_S.Department, rtnum=rtid)
    print (rtid)
    arec = DocxTemplate( gerpath +'/模板doc/REC.docx')

    arec.render(context)

    filepath = gerpath + '/record/放疗前讨论记录/' + pt_S.Names + '_放疗前讨论记录.docx'

    arec.save(filepath)

def first_mk(pt_S): #生成放疗首页


    ffile = Document()#'/home/zxf/病人登记/放疗前讨论记录/'+ pt_S.Names + '_放疗前讨论记录.docx')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    paragraph1 = ffile.add_paragraph('成都三六三医院')
    paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph1.size =Pt(26)
    # paragraph1.style = 'Title'
    paragraph2 = ffile.add_paragraph('放射治疗专科病例')
    paragraph2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph1.size = Pt(25)
    # paragraph2.style = 'Title'
    #  ffile.add_heading('成都三六三医院', 0)
    # ffile.add_heading('放射治疗专科病例', 0)
    """
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True  # 加粗
    p.add_run(' and some ')
    p.add_run('italic.').italic = True  # 倾斜
   
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    #document.add_paragraph(
    #    'first item in ordered list', style='ListNumber'
    #) """

    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')
    ffile.add_paragraph('')

    p = ffile.add_paragraph()
    p.add_run('                                                            病人姓名:    '+pt_S.Names)  # 加粗
    p.add_run('\n                                                            放   疗  号:    ' + pt_S.name)
    p.add_run('\n                                                            科        别:    ' + pt_S.Department) # 倾斜
    p.add_run('\n                                                            床        号:    ' + pt_S.bedNum)
    p.add_run('\n                                                            住   院   号:  ' + pt_S.id)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    #document.add_picture('monty-truth.png', width=Inches(1.25))
    """
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
   
    for item in recordset:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.qty)
        row_cells[1].text = str(item.id)
        row_cells[2].text = item.desc
    """
    ffile.add_page_break()

    ffile.save(gerpath +'/放疗前讨论记录/'+ pt_S.Names + '_放疗前讨论记录.docx')

