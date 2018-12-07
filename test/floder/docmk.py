
class doc_mk():

    def doc_mk(pt_S, beamNum, firmStyle, Tdose, date, diagnosisT Ldose, sectNum):


        import docx
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 引入文字的位置样式
        from docx.shared import Pt  # 设置字体
        import sys

        gerpath = str(sys.path[0])
        #
        #    需要传入:  病人信息  治疗野数^  固定方式^  医师  总治疗剂量^  date[] 特殊表现({{diagnosisT}})   需要UI支持三个量输入
        #   1.标题生成  done
        #   2.首次病程生成  done
        #   3.计算需要的中间病程数  ing...  计算结束日期OK , 中间记录数= 分割次数/5-2  计算每次日期. date[]
        #   4.生成自动的中间病程    循环实现  特殊表现代入与循环变量挂钩  5, 10, 15, 20, 25, 30, 35,特殊表现( {{ diagnosisT[] }} )
        #   5.文档变量代入  模板变量代入
        #   6.文档保存 文件名 病人姓名_病程.docx
        #   文档实现功能: 文档生成病程记录: 自动实现各个信息变量代入,格式代入,自动保存.


        gerpath = str(sys.path[0])
        file = docx.Document(gerpath + '/模板doc/text.docx')  # 新建word文档
        file2 = docx.Document(gerpath+'/模板doc/paragraph.docx')

        # info_get()返回 pt_S 为病人信息序列: ['id', 'RTid','Curesta', 'Names', 'sex', 'age', 'Department',
        #                                   'bedNum','dignose', 'phoneNum', 'bodyPart', 'startDate']
        #
        #  ['姓名:', pt_S.Names,       '性别:', pt_S.sex,           '年龄:', str(pt_S.age),
        #   '科室:', pt_S.Department,  '床号:', str(pt_S.bedNum),
        #   '诊断:', pt_S.dignose,     '电话:', str(pt_S.phoneNum), ,
        #   '治疗部位:', str(pt_S.bodyPart),    '开始日期:', str(pt_S.startDate)]
        #
        #   模板文档; 首次病程: X1doc.docx  段落文档: paragraph.docx
        #   需迭代变量:
        # {{startDate}}  --  date[0]^^^             date[]数组构建
        # {{dateN}}          date[n]
        # {{dateL}}          date[L]
        # {{pName}},         pt_S.Names
        # {{sex}},           pt_S.sex
        # {{age}}            pt_S.age
        # {{diagnose}}       pt_S.dignose
        # {{bodyPart}}       pt_S.bodyPart
        # {{confirm}}        firmStyle              固定方式传倒
        # {{beamNum}}        beamNum                野数传到
        # {{Dose}}           Tdose                  剂量推倒分割方式推倒 dose[]构建
        # {{DoseN}}          dose[n]
        # {{DoseL}}          Tdose/ Tdose+Ldose[]   分段模式构建总剂量计算
        # {{doctor}}    --   默认值
        # 内在变量sectNum 治疗分段改野次数
        # 函数传入列表 pt_S, beamNum, firmStyle, Tdose, date, diagnosisT Ldose, sectNum)

        # UI需要构建传入结构: 治疗野数^  固定方式^  总治疗剂量^  开始治疗日期^  分段数  [ Ldose[] ]
        # 野数输入框  固定方式下拉菜单  总剂量输入  开始日期输入框 分段数   根据分段数变化的LDose[]输入框数

        """
        #第一段
        paragraph1=file.add_paragraph('病程记录')
        paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph1.style='Title'

        #第二段
        paragraph2=file.add_paragraph()
        paragraph2.paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        L = 'text'
        run=paragraph2.add_run(L) #向段落中追加文字
        run.bold=True    #设置追加文字样式
        run.font.size=Pt(30)
        #run.style='Emphasis'

        # 第三段
        paragraph3=file.add_paragraph()
        paragraph3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph3.style = 'Normal'
        run=paragraph3.add_run('Tomorrow is my birthay.I am looking forword your coming')
        run.font.size=Pt(20)
        file.save('/home/zxf/病人登记/模板doc/text.docx')
        #file.save('/home/zxf/病人登记/模板doc/'+ L +'.docx')
        """
        docTex2 = [paragraph.text for paragraph in file2.paragraphs]
        for n in range():
            paragraphd = file.add_paragraph()
            paragraphd.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            ad0 = paragraphd.add_run(docTex2[0])
            ad0.font.size = Pt(11)

            paragraphn = file.add_paragraph()
            ad = paragraphn.add_run(docTex2[1])
            ad.font.size = Pt(11)

            paragraphs = file.add_paragraph()
            paragraphs.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            ad1 = paragraphs.add_run(docTex2[2])
            ad1.font.size = Pt(11)

        file.save('/home/zxf/病人登记/模板doc/text.docx')
