"""
from docx.enum.style import WD_STYLE_TYPE
from docx import *

document = Document()
styles = document.styles
para = document.add_paragraph()

#生成所有字符样式
for s in styles:
    if s.type == WD_STYLE_TYPE.CHARACTER:
        run = para.add_run("Character style is:  "+s.name+"\n")
        run.style = s

document.save('/home/zxf/病人登记/模板doc/character_style.docx')
"""

from docx.enum.style import WD_STYLE_TYPE
from docx import *

document = Document()
styles = document.styles

#生成所有段落样式
for s in styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH:
        document.add_paragraph('Paragraph style is : '+ s.name, style = s)

document.save('/home/zxf/病人登记/模板doc/para_style.docx')