

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #引入文字的位置样式
from docx.shared import Pt  #设置字体





file = docx.Document('/home/zxf/病人登记/模板doc/text.docx') #新建word文档
file2 = docx.Document('/home/zxf/病人登记/模板doc/paragraph.docx')
"""
#第一段
paragraph1=file.add_paragraph('病程记录')
paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
paragraph1.style='Title'

#第二段
paragraph2=file.add_paragraph()
paragraph2.paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
L = 'text'
run=paragraph2.add_run(L) #向段落中追加文字
run.bold=True    #设置追加文字样式
run.font.size=Pt(30)
#run.style='Emphasis'

# 第三段
paragraph3=file.add_paragraph()
paragraph3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
paragraph3.style = 'Normal'
run=paragraph3.add_run('Tomorrow is my birthay.I am looking forword your coming')
run.font.size=Pt(20)
file.save('/home/zxf/病人登记/模板doc/text.docx')
#file.save('/home/zxf/病人登记/模板doc/'+ L +'.docx')
"""
docTex2= [paragraph.text for paragraph in file2.paragraphs]
for n in range(4):

    paragraphd = file.add_paragraph()
    paragraphd.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ad0 = paragraphd.add_run(docTex2[0])
    ad0.font.size = Pt(11)

    paragraphn = file.add_paragraph()
    ad = paragraphn.add_run(docTex2[1])
    ad.font.size = Pt(11)

    paragraphs = file.add_paragraph()
    paragraphs.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    ad1 = paragraphs.add_run(docTex2[2])
    ad1.font.size = Pt(11)

file.save('/home/zxf/病人登记/模板doc/text.docx')
